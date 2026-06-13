#!/usr/bin/env python3
"""
Chenier Environmental Consulting
Figure 3 — Current Topographic Map (latest year only)

Reuses the proven discovery + rendering engine from topo_core_clean.py, but
selects ONLY the most recent topographic map year (e.g. 2024 US Topo) and
formats it to the Figure 3 template.

Files needed in the repo:
  streamlit_topo_current.py, topo_core_clean.py, north_arrow.jpeg,
  requirements.txt, packages.txt
"""

import io
import tempfile
import shutil
from pathlib import Path

import streamlit as st

import topo_core_clean as core

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NORTH_ARROW_PATH = 'north_arrow.jpeg'

# Page layout — matches the Figure 3 template (portrait, same as historic topo)
PAGE_W  = Inches(8.5);  PAGE_H  = Inches(11.0)
MAR_L   = Inches(0.7);  MAR_R   = Inches(0.25)
MAR_TOP = Inches(0.5);  MAR_BOT = Inches(0.5)
MAP_W   = Inches(7.5);  MAP_H   = Inches(8.5)


def _latest_year(item):
    """Best-effort integer year for an item, for picking the most recent map."""
    m = core.parse_meta(item)
    try:
        return int(float(m['year']))
    except (ValueError, TypeError):
        return -1


def build_doc_fig3(img_path, meta, project_no, out_buf, north_arrow):
    """Build the single-page Figure 3 document matching the template."""
    doc = Document()
    sec = doc.sections[0]
    sec.page_width  = PAGE_W;  sec.page_height = PAGE_H
    sec.left_margin = MAR_L;   sec.right_margin = MAR_R
    sec.top_margin  = MAR_TOP; sec.bottom_margin = MAR_BOT
    sec.header_distance = Inches(0); sec.footer_distance = Inches(0)
    for p in doc.paragraphs:
        p._element.getparent().remove(p._element)

    cw    = PAGE_W - MAR_L - MAR_R
    cap_w = Inches(5.3)
    key_w = Inches(2.25)
    year  = meta.get('year', '')

    # ── Blank-page guard (same approach as historic topo / site plan) ──────
    # Caption line: "USGS <Quad>, <State>". A multi-quad mosaic name can wrap;
    # shrink the map height to keep the page from overflowing.
    title_str = f"USGS {meta['quad']}, {meta['state']}"
    CHARS_PER_LINE = 74
    title_lines = max(1, -(-len(title_str) // CHARS_PER_LINE))
    extra_lines = max(0, title_lines - 1)
    map_h = MAP_H - Inches(0.24 * extra_lines)
    cap_pt = 10 if extra_lines else 11

    # ── Map image, framed ──────────────────────────────────────────────────
    mt = doc.add_table(1, 1)
    core.tbl_border(mt)
    mt.alignment = WD_TABLE_ALIGNMENT.LEFT
    tblPr = core.get_or_add(mt._tbl, 'w:tblPr')
    tblW = OxmlElement('w:tblW'); tblW.set(qn('w:w'), str(int(cw/914400*1440)))
    tblW.set(qn('w:type'), 'dxa'); tblPr.append(tblW)
    mc = mt.cell(0, 0); core.cell_w(mc, cw)
    tcPr = core.get_or_add(mc._tc, 'w:tcPr'); tcMar = OxmlElement('w:tcMar')
    for edge in ('top','left','bottom','right'):
        el = OxmlElement(f'w:{edge}'); el.set(qn('w:w'),'0'); el.set(qn('w:type'),'dxa')
        tcMar.append(el)
    tcPr.append(tcMar)
    mp = mc.paragraphs[0]; mp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    mp.paragraph_format.space_before = Pt(0); mp.paragraph_format.space_after = Pt(0)
    if Path(img_path).exists():
        mp.add_run().add_picture(img_path, width=MAP_W, height=map_h)

    # ── Caption row ──────────────────────────────────────────────────────────
    ct = doc.add_table(1, 2); ct.alignment = WD_TABLE_ALIGNMENT.LEFT
    lc = ct.cell(0,0); rc = ct.cell(0,1)
    core.cell_w(lc, cap_w); core.cell_w(rc, key_w)
    core.no_border(lc); core.no_border(rc)

    # Left: "USGS <Quad>, <State>" / ", 7.5-minutes; <scale>" / "Created: <year>"
    lp1 = lc.paragraphs[0]
    lp1.paragraph_format.space_before = Pt(0); lp1.paragraph_format.space_after = Pt(0)
    core.run(lp1, f"USGS {meta['quad']}, {meta['state']}", italic=True, pt=cap_pt)
    scale_txt = meta.get('scale', '') or '1:24,000'
    core.run(lp1, f", 7.5-minutes; {scale_txt}", pt=cap_pt)
    lp2 = lc.add_paragraph()
    lp2.paragraph_format.space_before = Pt(0); lp2.paragraph_format.space_after = Pt(0)
    core.run(lp2, f"Created: {year}", pt=cap_pt)

    # Right: KEY block
    rp1 = rc.paragraphs[0]
    rp1.paragraph_format.space_before = Pt(0); rp1.paragraph_format.space_after = Pt(0)
    core.run(rp1, 'KEY:', bold=True, pt=9)
    rp2 = rc.add_paragraph()
    rp2.paragraph_format.space_before = Pt(0); rp2.paragraph_format.space_after = Pt(0)
    if north_arrow and Path(north_arrow).exists():
        rp2.add_run().add_picture(str(north_arrow), height=Inches(0.28))
        core.run(rp2, '  ', pt=9)
    sym = rp2.add_run('━━  '); sym.font.color.rgb = RGBColor(255,0,0); sym.font.size = Pt(12)
    core.run(rp2, 'Subject Property', pt=9)

    # ── Footer ───────────────────────────────────────────────────────────────
    ftr = sec.footer
    for p in ftr.paragraphs:
        p._element.getparent().remove(p._element)
    ft = ftr.add_table(1, 3, width=Inches(7.55)); ft.alignment = WD_TABLE_ALIGNMENT.LEFT
    lc2 = ft.cell(0,0); mc2 = ft.cell(0,1); rc2 = ft.cell(0,2)
    core.cell_w(lc2, Inches(4.0)); core.cell_w(mc2, Inches(1.5)); core.cell_w(rc2, Inches(2.05))
    for c in (lc2,mc2,rc2): core.no_border(c)
    lf = lc2.paragraphs[0]
    r1 = lf.add_run(f'Figure 3:  {year} Topographic Map')
    r1.bold=True; r1.font.name='Segoe UI'; r1.font.size=Pt(14); r1.font.all_caps=True
    lf2 = lc2.add_paragraph()
    r2 = lf2.add_run(f'Project No. {project_no}')
    r2.bold=True; r2.font.name='Segoe UI'; r2.font.size=Pt(10)
    rf = rc2.paragraphs[0]; rf.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    rf.paragraph_format.space_before = Pt(6)
    cr = rf.add_run('Chenier Environmental Consulting, LLC')
    cr.font.name='Segoe UI'; cr.font.size=Pt(10)
    tP = core.get_or_add(ft._tbl,'w:tblPr'); brd = OxmlElement('w:tblBorders')
    tp = OxmlElement('w:top'); tp.set(qn('w:val'),'single'); tp.set(qn('w:sz'),'4')
    tp.set(qn('w:color'),'auto'); brd.append(tp); tP.append(brd)

    doc.save(out_buf)


# ── Streamlit UI ─────────────────────────────────────────────────────────────────
st.set_page_config(page_title="Chenier — Current Topo Map", page_icon="🗻", layout="centered")
st.title("🗻 Current Topographic Map Generator")
st.caption("Chenier Environmental Consulting, LLC — Figure 3 (latest USGS topo)")
st.markdown("---")

uploaded = st.file_uploader("**1. Upload site boundary (KMZ or KML)**", type=['kmz','kml'])
project_no = st.text_input("**2. Project number**", placeholder="e.g. 26-014")

c1, c2 = st.columns(2)
with c1:
    zoom_buffer = st.slider("**Zoom out**", 0.10, 1.00, 0.35, 0.05,
                            help="Higher = more area shown around the site")
with c2:
    collar_trim = st.slider("**Collar trim**", 0.000, 0.010, 0.000, 0.0005,
                            format="%.4f",
                            help="How much white quad border to crop when "
                                 "mosaicking. Increase if white edges show; "
                                 "decrease if real map content gets cut off.")

generate = st.button("⚡ Generate Current Topographic Map", type="primary", use_container_width=True)

if generate:
    if not uploaded:
        st.error("Please upload a KMZ or KML file.")
    elif not project_no.strip():
        st.error("Please enter a project number.")
    else:
        work = Path(tempfile.mkdtemp(prefix="chenier_fig3_"))
        try:
            with st.status("Generating current topo map...", expanded=True) as status:
                st.write("Reading boundary file...")
                fb = uploaded.read()
                site_geom, bounds = core.parse_kmz_bytes(fb, uploaded.name)

                st.write("Querying USGS National Map...")
                items = core.query_tnm(bounds)
                if not items:
                    status.update(label="No maps found", state="error")
                    st.error("No topographic maps found for this area.")
                    st.stop()

                if getattr(core, 'US_TOPO_WARNING', None):
                    st.warning("⚠ " + core.US_TOPO_WARNING +
                               "  (Figure 3 needs the most recent map, so this "
                               "matters — re-run when USGS US Topo responds.)")

                # Pick the most recent year available
                years = [_latest_year(i) for i in items]
                max_year = max(years)
                latest_items = [i for i, y in zip(items, years) if y == max_year]
                st.write(f"Most recent map year: {max_year} "
                         f"({len(latest_items)} quad(s) at this year)")

                # Group only the latest-year items (handles a site spanning quads)
                groups = core.group_adjacent_quads(latest_items, bounds)
                grp = groups[0]  # the group covering the site

                cb = core.clip_bounds(bounds, buffer_factor=zoom_buffer)
                tmp = work / "_tmp"; tmp.mkdir(exist_ok=True)

                metas = [core.parse_meta(i) for i in grp]
                if len(grp) == 1:
                    meta = metas[0]
                else:
                    meta = metas[0].copy()
                    meta['quad'] = ' / '.join(m['quad'] for m in metas)

                st.write(f"Downloading {len(grp)} quad(s)...")
                tif_paths = []
                for k, item in enumerate(grp):
                    tif = tmp / f"q{k}.tif"
                    core.download_tif(item['_tif'], tif)
                    tif_paths.append(tif)

                st.write("Rendering map...")
                jpg = tmp / "fig3.jpg"
                if not core.render_group(tif_paths, site_geom, cb, jpg,
                                         items=grp, collar_shrink=collar_trim):
                    status.update(label="Render failed", state="error")
                    st.error("Could not render the map.")
                    st.stop()

                st.write("Building Word document...")
                na = NORTH_ARROW_PATH if Path(NORTH_ARROW_PATH).exists() else None
                if not na:
                    st.write("(north_arrow.jpeg not found — north arrow omitted)")

                buf = io.BytesIO()
                build_doc_fig3(str(jpg), meta, project_no.strip(), buf, na)
                buf.seek(0)
                docx_bytes = buf.getvalue()
                preview = jpg.read_bytes()

                status.update(label=f"Done — {max_year} map", state="complete", expanded=False)

            st.success(f"Generated Figure 3 — {max_year} Topographic Map.")
            safe = project_no.strip().replace('/','_').replace('\\','_')
            st.download_button(
                "⬇ Download Word Document",
                data=docx_bytes,
                file_name=f"{safe}_Fig_3_Topographic_Map.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                type="primary", use_container_width=True)
            st.image(preview, caption=f"Preview — {max_year}", use_container_width=True)

        except Exception as e:
            st.error(f"Error: {e}")
            import traceback; st.code(traceback.format_exc())
        finally:
            shutil.rmtree(work, ignore_errors=True)

st.markdown("---")
st.caption("Map data: USGS The National Map")
