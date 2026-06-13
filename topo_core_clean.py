#!/usr/bin/env python3
"""
Chenier Environmental Consulting
Phase I ESA - Historic Topographic Figure Generator
v3 - Edition info, frame fix, adjacent quad mosaicking
"""

import os, sys, io, zipfile, shutil, subprocess, re, json
from pathlib import Path
from datetime import datetime
from collections import defaultdict

REQUIRED = {'requests':'requests','rasterio':'rasterio','pyproj':'pyproj',
            'PIL':'Pillow','docx':'python-docx','shapely':'shapely','lxml':'lxml'}

def ensure_packages():
    import importlib
    missing = [pkg for mod,pkg in REQUIRED.items() if not importlib.util.find_spec(mod)]
    if missing:
        print(f"Installing: {', '.join(missing)} ...")
        subprocess.check_call([sys.executable,'-m','pip','install','--quiet']+missing)
        print("Done.\n")

# ensure_packages()  # disabled for Streamlit

import requests, numpy as np
import rasterio
from rasterio.warp import Resampling
from rasterio.windows import from_bounds as win_from_bounds
from rasterio.merge import merge as rasterio_merge
from rasterio.transform import from_bounds as transform_from_bounds
from pyproj import Transformer
from PIL import Image, ImageDraw
from shapely.geometry import Polygon
from shapely.ops import unary_union
from lxml import etree
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

TNM_API        = 'https://tnmaccess.nationalmap.gov/api/v1/products'
# USGS endpoints sometimes 500 or throttle the default python-requests
# User-Agent. Sending a normal browser-like UA avoids that. This is the
# main behavioral difference between a working browser request and the app.
USGS_HEADERS   = {
    'User-Agent': ('Mozilla/5.0 (Windows NT 10.0; Win64; x64) '
                   'AppleWebKit/537.36 (KHTML, like Gecko) '
                   'Chrome/124.0 Safari/537.36 ChenierTopoTool/1.0'),
    'Accept': 'application/json',
}
OVERLAY_URL    = 'https://raw.githubusercontent.com/Brad-Chenier/chenier-data/main/USGS_Topo_Overlay_Data_new.geojson'
BUFFER_FACTOR  = 0.35
MIN_BUFFER     = 0.04
# How far (in degrees) to crop inward from each quad's declared map-sheet
# boundingBox before mosaicking. The TNM boundingBox is the full sheet
# including the white collar/margin, so we must crop past it to the neat-line
# (actual map content). ~0.0045 deg ≈ 0.5 km, which clears a 7.5' quad collar.
COLLAR_SHRINK_DEG = 0.0045
SITE_COLOR     = (255, 0, 0)  # red boundary
IMG_W_PX       = 1125   # 7.5" x 150 DPI
IMG_H_PX       = 1275   # 8.5" x 150 DPI
JPEG_QUALITY   = 80
PAGE_W         = Inches(8.5)
PAGE_H         = Inches(11.0)
MAR_L          = Inches(0.7)
MAR_R          = Inches(0.25)
MAR_TOP        = Inches(0.5)
MAR_BOT        = Inches(0.5)
MAP_IMG_W      = Inches(7.5)              # map width
MAP_IMG_H      = Inches(8.5)              # map height
SCRIPT_DIR     = Path(__file__).parent


# ── KMZ ──────────────────────────────────────────────────────────────────────
def parse_kmz(path):
    with zipfile.ZipFile(path) as z:
        names = [n for n in z.namelist() if n.lower().endswith('.kml')]
        kml   = z.read('doc.kml' if 'doc.kml' in names else names[0])
    root  = etree.fromstring(kml)
    polys = []
    for el in root.iter('{http://www.opengis.net/kml/2.2}coordinates'):
        pts = []
        for tok in el.text.strip().split():
            p = tok.split(',')
            if len(p) >= 2:
                try: pts.append((float(p[0]), float(p[1])))
                except: pass
        if len(pts) >= 3: polys.append(Polygon(pts))
    if not polys: raise ValueError("No polygon found in KMZ")
    g = unary_union(polys)
    return g, g.bounds


def parse_kmz_bytes(file_bytes, filename):
    """Parse KMZ or KML from in-memory bytes (for web uploads).
    Returns (geom, bounds)."""
    if filename.lower().endswith('.kmz'):
        with zipfile.ZipFile(io.BytesIO(file_bytes)) as z:
            names = [n for n in z.namelist() if n.lower().endswith('.kml')]
            if not names:
                raise ValueError("No KML found inside KMZ")
            kml = z.read('doc.kml' if 'doc.kml' in names else names[0])
    else:
        kml = file_bytes
    root  = etree.fromstring(kml)
    polys = []
    for el in root.iter('{http://www.opengis.net/kml/2.2}coordinates'):
        pts = []
        for tok in el.text.strip().split():
            p = tok.split(',')
            if len(p) >= 2:
                try: pts.append((float(p[0]), float(p[1])))
                except: pass
        if len(pts) >= 3: polys.append(Polygon(pts))
    if not polys: raise ValueError("No polygon found in KMZ/KML")
    g = unary_union(polys)
    return g, g.bounds


# ── Overlay GeoJSON (edition metadata) ───────────────────────────────────────
_overlay_cache = None

def load_overlay():
    global _overlay_cache
    if _overlay_cache is not None:
        return _overlay_cache
    print("  Loading USGS topo overlay data (edition info)...")
    try:
        r = requests.get(OVERLAY_URL, timeout=30)
        r.raise_for_status()
        data = r.json()
        _overlay_cache = {}
        for feat in data.get('features', []):
            p   = feat.get('properties', {})
            sid = str(p.get('scan_id', ''))
            if sid:
                _overlay_cache[sid] = p
        print(f"  Loaded {len(_overlay_cache)} overlay records")
    except Exception as e:
        print(f"  Warning: could not load overlay data ({e})")
        _overlay_cache = {}
    return _overlay_cache



# Set to a message string when the most recent query_tnm() call could not
# retrieve the US Topo (recent, 2009-present) dataset. The UI checks this so a
# silent gap in recent maps can be surfaced to the user.
US_TOPO_WARNING = None


def _fetch_json_with_retry(url, params=None, headers=None, timeout=45,
                           max_attempts=6, label='USGS'):
    """GET a URL and parse JSON, retrying on transient failures.

    USGS endpoints are intermittently unstable — they return 500s, drop
    connections mid-response, or return non-JSON error bodies. This retries
    with backoff and returns the parsed JSON dict. Raises RuntimeError if all
    attempts fail (so callers can choose to catch or propagate).
    """
    import time as _time
    last_err = None
    r = None
    for attempt in range(max_attempts):
        try:
            r = requests.get(url, params=params, headers=headers, timeout=timeout)
            r.raise_for_status()
            return r.json()
        except ValueError as e:           # 200 OK but body isn't JSON
            last_err = e
            print(f"    {label} attempt {attempt+1}/{max_attempts}: non-JSON response; retrying...")
            _time.sleep(min(3 * (attempt + 1), 12))
        except requests.RequestException as e:   # 500, timeout, dropped conn, etc.
            last_err = e
            print(f"    {label} attempt {attempt+1}/{max_attempts}: request failed ({e}); retrying...")
            _time.sleep(min(3 * (attempt + 1), 12))
    snippet = ''
    try:
        snippet = (r.text or '').strip()[:300]
    except Exception:
        pass
    raise RuntimeError(
        f"The USGS National Map API ({label}) failed after {max_attempts} "
        "attempts. Their service is intermittently unstable, so this usually "
        "clears up within a few minutes — please wait a moment and try again."
        f"\n\nLast response from USGS:\n{snippet!r}\n(Underlying error: {last_err})")


def query_tnm(bounds):
    x0,y0,x1,y1 = bounds
    # Expand query bbox to match the clip window (BUFFER_FACTOR + margin)
    # so that any neighbor quads visible in the rendered figure are fetched
    bx = max((x1-x0) * BUFFER_FACTOR, MIN_BUFFER) + 0.02
    by = max((y1-y0) * BUFFER_FACTOR, MIN_BUFFER) + 0.02
    qx0 = x0-bx; qy0 = y0-by; qx1 = x1+bx; qy1 = y1+by
    # Round to 6 decimal places (~0.1 m). Full float precision produces
    # 17-digit coordinate strings that the USGS TNM API's parameter parser
    # rejects as a BadRequest, so the bbox must be sent at sane precision.
    qx0 = round(qx0, 6); qy0 = round(qy0, 6)
    qx1 = round(qx1, 6); qy1 = round(qy1, 6)
    print("  Querying USGS TNM...")
    global US_TOPO_WARNING
    US_TOPO_WARNING = None
    import time as _time

    # Historical Topographic Maps (HTMC) — required. Propagates on failure.
    d = _fetch_json_with_retry(
        TNM_API,
        params={'datasets':'Historical Topographic Maps',
                'bbox':f'{qx0},{qy0},{qx1},{qy1}',
                'max':200, 'outputFormat':'JSON'},
        headers=USGS_HEADERS, timeout=45, label='Historical Topo')
    items = d.get('items', [])
    print(f"  Found {len(items)} HTMC maps (total: {d.get('total','?')})")
    valid = []
    for i in items:
        tif = i.get('urls',{}).get('GeoTIFF') or i.get('geotiffURL')
        if tif:
            i['_tif'] = tif
            i['_series'] = 'HTMC'
            valid.append(i)
    print(f"  {len(valid)} have GeoTIFF downloads")

    # Also query US Topo (2009-present) — the recent maps. Optional: if it
    # fails after retries we keep the HTMC results but record a warning so the
    # UI can tell the user recent maps may be missing (instead of silently
    # producing a document that looks complete).
    try:
        _time.sleep(0.5)  # brief spacing to avoid burst-throttling from USGS
        d2 = _fetch_json_with_retry(
            TNM_API,
            params={'datasets':'US Topo',
                    'bbox':f'{qx0},{qy0},{qx1},{qy1}',
                    'max':200, 'outputFormat':'JSON'},
            headers=USGS_HEADERS, timeout=30, label='US Topo')
        us_items = d2.get('items', [])
        print(f"  Found {len(us_items)} US Topo maps (total: {d2.get('total','?')})")
        added = 0
        for i in us_items:
            # US Topo GeoTIFFs are not in urls.GeoTIFF — derive from the PDF downloadURL
            # PDF:    .../USTopo/PDF/LA/LA_QuadName_20240101_TM_geo.pdf
            # GeoTIF: .../USTopo/GeoTIFF/LA/LA_QuadName_20240101_TM_geo.tif
            pdf_url = i.get('downloadURL') or i.get('urls', {}).get('GeoPDF', '')
            tif = None
            if pdf_url and 'USTopo' in pdf_url:
                tif = pdf_url.replace('/USTopo/PDF/', '/USTopo/GeoTIFF/') \
                             .replace('_TM_geo.pdf', '_TM_geo.tif') \
                             .replace('_geo.pdf', '_geo.tif')
            # Also check if GeoTIFF is directly provided
            if not tif:
                tif = i.get('urls', {}).get('GeoTIFF') or i.get('geotiffURL')
            if tif:
                i['_tif'] = tif
                i['_series'] = 'US Topo'
                valid.append(i)
                added += 1
        print(f"  {added} US Topo maps added  ({len(valid)} total)")
    except Exception as e:
        US_TOPO_WARNING = (
            "Recent US Topo maps (2009-present) could NOT be retrieved this "
            "run because the USGS US Topo service did not respond. The "
            "document below contains the older historical maps only. This is "
            "usually a temporary USGS issue — re-run in a few minutes to "
            "include the recent maps.")
        print(f"  Warning: US Topo query failed after retries ({e})")

    valid.sort(key=lambda i: i.get('publicationDate',''), reverse=True)

    # Supplement with HTMC revised editions via ustOverlay ArcGIS service
    # This returns only maps covering this specific location with full edition metadata
    try:
        overlay_url = (
            'https://ngmdb-tiles.usgs.gov/arcgis/rest/services/topoview/ustOverlay/MapServer/0/query'
            f'?where=date_on_map+>=+1880'
            f'&geometry={qx0},{qy0},{qx1},{qy1}'
            f'&geometryType=esriGeometryEnvelope'
            f'&spatialRel=esriSpatialRelIntersects'
            f'&outFields=scan_id,date_on_map,imprint_year,series,map_name,map_scale,primary_state'
            f'&returnGeometry=false'
            f'&resultRecordCount=200'
            f'&f=json'
        )
        _time.sleep(0.5)  # brief spacing to avoid burst-throttling
        ov_r = requests.get(overlay_url, headers=USGS_HEADERS, timeout=20)
        ov_r.raise_for_status()
        ov_data = ov_r.json()
        ov_features = ov_data.get('features', [])
        print(f"  ustOverlay: {len(ov_features)} map editions at this location")

        # Build set of scan_ids already in TNM results
        tnm_scans = set()
        for item in valid:
            m2 = re.search(r'_([0-9]{6,7})_[0-9]{4}_', item.get('_tif',''))
            if m2: tnm_scans.add(m2.group(1))

        extra = []
        for feat in ov_features:
            a = feat.get('attributes', {})
            sid         = str(a.get('scan_id', '') or '').strip()
            date_on_map = str(a.get('date_on_map', '') or '').strip()
            imprint_yr  = str(a.get('imprint_year', '') or '').strip()
            series      = str(a.get('series', '') or '').strip()
            map_name    = str(a.get('map_name', '') or '').strip()
            map_scale   = str(a.get('map_scale', '') or '').strip()
            state       = str(a.get('map_state', '') or a.get('primary_state', '') or '').strip().upper()

            if not sid or sid in tnm_scans:
                continue
            if not map_name or not map_scale or not date_on_map:
                continue

            name_clean = map_name.replace(' ','_').replace('/','_')
            try:
                scale_int = int(float(map_scale))
            except:
                continue

            tif_url = (f"https://prd-tnm.s3.amazonaws.com/StagedProducts/Maps/"
                       f"HistoricalTopo/GeoTIFF/{state}/"
                       f"{state}_{name_clean}_{sid}_{date_on_map}_{scale_int}_geo.tif")

            extra.append({
                'title':           f"USGS 1:{scale_int:,}-scale Quadrangle for {map_name}, {state} {date_on_map}",
                'publicationDate': f"{date_on_map}-01-01",
                '_tif':            tif_url,
                '_from_overlay':   True,
                'boundingBox':     None,
                '_imprint_yr':     imprint_yr,
                '_series':         series,
            })

        if extra:
            print(f"  + {len(extra)} HTMC revised editions added")
            valid.extend(extra)
        else:
            print(f"  No additional HTMC editions found")

    except Exception as e:
        print(f"  Warning: ustOverlay query failed ({e})")

    return valid


def get_edition_str(tif_url, year, imprint_override=None, series_override=None):
    """Return edition string like ' (HTMC, 1994 ed.)' or '' if none."""
    # For overlay-sourced items, use pre-stored values
    if imprint_override:
        try:
            imp_norm = str(int(float(imprint_override)))
            yr_norm  = str(int(float(year))) if year else year
        except:
            imp_norm = imprint_override
            yr_norm  = year
        if imp_norm and imp_norm != yr_norm:
            series = series_override or ''
            parts  = [series, imp_norm + ' ed.'] if series else [imp_norm + ' ed.']
            return ' (' + ', '.join(parts) + ')'
        return ''
    # For TNM items, look up by scan_id from filename
    m = re.search(r'_(\d{6,7})_\d{4}_', tif_url or '')
    if not m: return ''
    overlay = load_overlay()
    raw_sid = m.group(1)
    props   = overlay.get(raw_sid) or overlay.get(str(int(raw_sid)), {}) or {}
    raw_imp = str(props.get('imprint_year', '') or props.get('imprint_ye', '') or '').strip()
    series  = str(props.get('series', '') or '').strip()
    try:
        imprint = str(int(float(raw_imp))) if raw_imp else ''
        yr_norm = str(int(float(year))) if year else year
    except:
        imprint = raw_imp
        yr_norm = year
    if imprint and imprint != yr_norm:
        parts = [series, imprint + ' ed.'] if series else [imprint + ' ed.']
        return ' (' + ', '.join(parts) + ')'
    return ''


def parse_meta(item):
    t = item.get('title', '')
    series = item.get('_series', '')

    # US Topo titles: "US Topo 7.5-minute map for QuadName, ST YYYY"
    # HTMC titles:    "USGS 1:24,000-scale Quadrangle for QuadName, ST YYYY"
    m_ustopo = re.search(r'(?:US Topo.*?map for|7\.5-minute map for)\s+(.+?),\s+([A-Z]{2})\s+(\d{4})', t)
    m_htmc   = re.search(r'Quadrangle for (.+?),?\s+(\d{4,6})', t)
    s        = re.search(r'1:([\d,]+)', t)

    quad = year = ''; state = 'Louisiana'

    if m_ustopo:
        quad  = m_ustopo.group(1).strip()
        state = m_ustopo.group(2).strip()
        year  = m_ustopo.group(3).strip()
        # US Topo is always 1:24,000
        if not s:
            scale = '1:24,000'
        else:
            scale = '1:' + s.group(1)
    elif series == 'US Topo':
        # Title has no year (e.g. "US Topo 7.5-minute map for Bee Bayou, LA")
        # Extract quad name and fall back to publicationDate for year
        m_name = re.search(r'(?:US Topo.*?map for|7\.5-minute map for)\s+(.+?),\s+([A-Z]{2})', t)
        if m_name:
            quad  = m_name.group(1).strip()
            state = m_name.group(2).strip()
        pub = item.get('publicationDate', '')
        year = pub[:4] if pub else ''
        scale = '1:24,000'
    elif m_htmc:
        parts = [x.strip() for x in m_htmc.group(1).split(',')]
        quad  = parts[0]
        state = parts[1] if len(parts) > 1 else 'Louisiana'
        try:
            year = str(int(float(m_htmc.group(2))))
        except:
            year = m_htmc.group(2)
        scale = ('1:' + s.group(1)) if s else ''
    else:
        scale = ('1:' + s.group(1)) if s else ''
    # For overlay-sourced items, use pre-stored edition info
    imp_override    = item.get('_imprint_yr','')
    series_override = item.get('_series','')
    # US Topo maps have no imprint year — label them as "(US Topo)" in the title block
    if series_override == 'US Topo':
        edition = ' (US Topo)'
    else:
        edition = get_edition_str(item.get('_tif',''), year,
                                  imp_override, series_override)
    return {'quad':quad or t, 'year':year, 'state':state,
            'scale':scale, 'edition':edition}


def scale_year_key(item):
    """Key for grouping spatially adjacent quads: (year, scale_number, imprint_year).
    Including imprint_year ensures different HTMC editions are separate figures."""
    t = item.get('title', '')

    # For US Topo items, year may only be in publicationDate, and scale is always 24000
    if item.get('_series') == 'US Topo':
        # Try title first (some have year at end), fall back to publicationDate
        m = re.search(r'\b(\d{4})\b', t)
        pub = item.get('publicationDate', '')
        base_year = m.group(1) if m else (pub[:4] if pub else '')
        return (base_year, '24000', base_year)

    m = re.search(r'(\d{4})$', t)
    s = re.search(r'1:([\d,]+)', t)
    base_year  = m.group(1) if m else ''
    scale_num  = s.group(1).replace(',','') if s else ''

    # For overlay-sourced items, use pre-stored imprint year directly
    if item.get('_from_overlay'):
        raw_imp = str(item.get('_imprint_yr', '') or '').strip()
        try:
            imprint_yr = str(int(float(raw_imp))) if raw_imp else ''
            base_yr_n  = str(int(float(base_year))) if base_year else base_year
        except:
            imprint_yr = raw_imp
            base_yr_n  = base_year
        edition_key = imprint_yr if imprint_yr and imprint_yr != base_yr_n else base_yr_n
        return (base_year, scale_num, edition_key)

    # For TNM items, look up imprint year from overlay data by scan_id
    tif_url    = item.get('_tif', '') or item.get('urls', {}).get('GeoTIFF', '')
    id_match   = re.search(r'_(\d{6,7})_\d{4}_', tif_url)
    imprint_yr = ''
    if id_match:
        overlay  = load_overlay()
        raw_sid  = id_match.group(1)
        try:
            norm_sid = str(int(raw_sid))
        except:
            norm_sid = raw_sid
        props      = overlay.get(norm_sid, {})
        raw_imp    = str(props.get('imprint_year', '') or props.get('imprint_ye', '') or '').strip()
        try:
            imprint_yr = str(int(float(raw_imp))) if raw_imp else ''
            base_yr_n  = str(int(float(base_year))) if base_year else base_year
        except:
            imprint_yr = raw_imp
            base_yr_n  = base_year
    else:
        base_yr_n = base_year
    # If imprint year differs from base year, it's a revised edition
    edition_key = imprint_yr if imprint_yr and imprint_yr != base_yr_n else base_yr_n
    return (base_year, scale_num, edition_key)


def group_adjacent_quads(items, site_bounds):
    """
    For each (year+scale) group, determine whether the clip window (the area
    that will actually be rendered) extends past a quad edge into the collar.
    If so, mosaic with the geographically adjacent quad on that side.
    """
    from shapely.geometry import box as sbox

    sx0, sy0, sx1, sy1 = site_bounds
    site_box = sbox(sx0, sy0, sx1, sy1)
    site_cx  = (sx0 + sx1) / 2
    site_cy  = (sy0 + sy1) / 2

    # Compute the clip window — same formula as clip_bounds()
    bx = max((sx1 - sx0) * BUFFER_FACTOR, MIN_BUFFER)
    by = max((sy1 - sy0) * BUFFER_FACTOR, MIN_BUFFER)
    cx0 = sx0 - bx;  cy0 = sy0 - by
    cx1 = sx1 + bx;  cy1 = sy1 + by

    # Group by year+scale
    groups = defaultdict(list)
    for item in items:
        key = scale_year_key(item)
        groups[key].append(item)

    result = []
    print(f'  Scale/year groups:')
    for key, grp in sorted(groups.items()):
        print(f'    {key}: {len(grp)} item(s) - {[i.get("title","")[:40] for i in grp]}')
    for key, grp in groups.items():
        # Attach shapely box to each item that has a boundingBox
        for item in grp:
            bb = item.get('boundingBox', {})
            if bb:
                item['_qbox'] = sbox(bb['minX'], bb['minY'],
                                     bb['maxX'], bb['maxY'])

        # Items with no boundingBox (overlay-supplemented) go straight to result
        no_bbox_items = [i for i in grp if '_qbox' not in i]
        bbox_items    = [i for i in grp if '_qbox' in i]

        # Each no-bbox item becomes its own figure
        for item in no_bbox_items:
            result.append([item])

        if not bbox_items:
            continue

        # Find primary quad: the one that contains the site centroid
        containing = [i for i in bbox_items
                      if i['_qbox'].contains(sbox(site_cx-0.001, site_cy-0.001,
                                                  site_cx+0.001, site_cy+0.001))]
        if not containing:
            containing = [i for i in bbox_items
                          if i['_qbox'].intersects(site_box)]
        if not containing:
            continue

        # Primary = largest intersection with site
        primary = max(containing,
                      key=lambda i: i['_qbox'].intersection(site_box).area)

        bb  = primary.get('boundingBox', {})
        qx0 = bb['minX']; qy0 = bb['minY']
        qx1 = bb['maxX']; qy1 = bb['maxY']
        qw  = qx1 - qx0;  qh  = qy1 - qy0

        # Trigger mosaic if the clip window extends past a quad edge
        # (i.e. the collar would be visible in the rendered figure)
        need_east  = cx1 > qx1
        need_west  = cx0 < qx0
        need_north = cy1 > qy1
        need_south = cy0 < qy0

        if not any([need_east, need_west, need_north, need_south]):
            # Clip window is entirely inside this quad — no collar visible,
            # no need for other quads from this year/scale group
            result.append([primary])
            continue

        # Clip window crosses a quad edge — iteratively expand mosaic group
        # by adding any candidate whose bbox intersects the clip window AND
        # shares an edge or corner with any quad already in the mosaic group.
        mosaic_group = [primary]
        tol = max(qw, qh) * 0.02   # 2% tolerance for shared edge detection
        clip_box = sbox(cx0, cy0, cx1, cy1)

        changed = True
        while changed:
            changed = False
            for candidate in bbox_items:
                if candidate in mosaic_group: continue
                if '_qbox' not in candidate: continue
                # Must overlap the clip window to be worth including
                if not clip_box.intersects(candidate['_qbox']):
                    continue
                # Must share an edge or corner with at least one mosaic member
                nb  = candidate.get('boundingBox', {})
                nx0 = nb['minX']; ny0 = nb['minY']
                nx1 = nb['maxX']; ny1 = nb['maxY']
                for member in mosaic_group:
                    mb  = member.get('boundingBox', {})
                    mx0 = mb['minX']; my0 = mb['minY']
                    mx1 = mb['maxX']; my1 = mb['maxY']
                    mtol = max(mx1-mx0, my1-my0) * 0.02
                    # Edge-adjacent (cardinal) or corner-adjacent (diagonal)
                    touches_x = abs(nx1-mx0)<mtol or abs(nx0-mx1)<mtol or \
                                 (nx0<mx1+mtol and nx1>mx0-mtol)
                    touches_y = abs(ny1-my0)<mtol or abs(ny0-my1)<mtol or \
                                 (ny0<my1+mtol and ny1>my0-mtol)
                    shares_edge   = (abs(nx1-mx0)<mtol or abs(nx0-mx1)<mtol) and touches_y
                    shares_edge  |= (abs(ny1-my0)<mtol or abs(ny0-my1)<mtol) and touches_x
                    shares_corner = (abs(nx1-mx0)<mtol or abs(nx0-mx1)<mtol) and \
                                    (abs(ny1-my0)<mtol or abs(ny0-my1)<mtol)
                    if shares_edge or shares_corner:
                        mosaic_group.append(candidate)
                        changed = True
                        break

        result.append(mosaic_group)

        # Drop any remaining quads — they don't overlap the clip window
        # and are only present because the TNM query bbox was expanded

    # Sort newest first
    result.sort(key=lambda g: g[0].get('publicationDate',''), reverse=True)
    return result


def clip_bounds(b, buffer_factor=None):
    x0,y0,x1,y1 = b
    bf = BUFFER_FACTOR if buffer_factor is None else buffer_factor
    bx = max((x1-x0)*bf, MIN_BUFFER)
    by = max((y1-y0)*bf, MIN_BUFFER)
    return (x0-bx, y0-by, x1+bx, y1+by)


def download_tif(url, dest):
    print("    Downloading...", end='', flush=True)
    with requests.get(url, stream=True, headers=USGS_HEADERS, timeout=180) as r:
        r.raise_for_status()
        total = int(r.headers.get('content-length', 0))
        done  = 0
        with open(dest, 'wb') as f:
            for chunk in r.iter_content(524288):
                f.write(chunk); done += len(chunk)
                if total:
                    print(f"\r    Downloading {done/total*100:.0f}% ({done//1048576} MB)", end='', flush=True)
    print()


def render_group(tif_paths, site_geom, clip_wgs84, out_jpg, items=None,
                 collar_shrink=None):
    """
    Render one or more GeoTIFFs (adjacent quads) into a single clipped JPEG.
    If multiple TIFs, mosaics them first then clips.
    collar_shrink: degrees to crop inward from each quad bbox (collar removal).
    """
    if len(tif_paths) == 1:
        return render_single(tif_paths[0], site_geom, clip_wgs84, out_jpg)
    else:
        return render_mosaic(tif_paths, site_geom, clip_wgs84, out_jpg,
                             items=items, collar_shrink=collar_shrink)


def render_single(tif_path, site_geom, clip_wgs84, out_jpg):
    with rasterio.open(tif_path) as src:
        return _clip_and_save(src, None, site_geom, clip_wgs84, out_jpg)



def render_mosaic(tif_paths, site_geom, clip_wgs84, out_jpg, items=None,
                  collar_shrink=None):
    """Mosaic adjacent quads, clipping each to its declared bbox before merging."""
    shrink_deg = COLLAR_SHRINK_DEG if collar_shrink is None else collar_shrink
    print(f"    Mosaicking {len(tif_paths)} adjacent quads... (collar shrink {shrink_deg} deg)")
    datasets       = []
    reproj_paths   = []
    collar_paths   = []
    tmp_mosaic     = tif_paths[0].parent / '_mosaic_tmp.tif'
    try:
        for k, p in enumerate(tif_paths):
            if not p.exists():
                continue
            ds = rasterio.open(p)

            # Clip to declared bounding box to remove collar
            # Use boundingBox from TNM item if available
            quad_bbox = None
            if items and k < len(items):
                bb = items[k].get('boundingBox', {})
                if bb:
                    quad_bbox = (bb['minX'], bb['minY'], bb['maxX'], bb['maxY'])

            if quad_bbox:
                # Reproject bbox to raster CRS and clip
                t_fwd = Transformer.from_crs('EPSG:4326', ds.crs, always_xy=True)
                qx0, qy0 = t_fwd.transform(quad_bbox[0], quad_bbox[1])
                qx1, qy1 = t_fwd.transform(quad_bbox[2], quad_bbox[3])
                # Shrink inward to crop past the collar to the neat-line.
                # The TNM boundingBox is the full sheet (with white collar), so
                # this must be large enough to reach actual map content.
                shrink = shrink_deg  # degrees (tunable per call)
                t2     = Transformer.from_crs('EPSG:4326', ds.crs, always_xy=True)
                sx0, sy0 = t2.transform(quad_bbox[0]+shrink, quad_bbox[1]+shrink)
                sx1, sy1 = t2.transform(quad_bbox[2]-shrink, quad_bbox[3]-shrink)
                clip_win = win_from_bounds(
                    min(sx0,sx1), min(sy0,sy1), max(sx0,sx1), max(sy0,sy1),
                    transform=ds.transform)
                clip_win = clip_win.intersection(
                    rasterio.windows.Window(0, 0, ds.width, ds.height))
                if clip_win.width > 0 and clip_win.height > 0:
                    clipped_data      = ds.read(window=clip_win)
                    clipped_transform = ds.window_transform(clip_win)
                    collar_path = p.parent / f'_clipped_{k}.tif'
                    profile = ds.profile.copy()
                    profile.update({
                        'driver': 'GTiff', 'compress': 'deflate',
                        'photometric': 'RGB' if ds.count >= 3 else 'MINISBLACK',
                        'height': clipped_data.shape[1],
                        'width':  clipped_data.shape[2],
                        'transform': clipped_transform,
                    })
                    for key in ('COMPRESS','PHOTOMETRIC','PREDICTOR','TILED',
                                'BLOCKXSIZE','BLOCKYSIZE','INTERLEAVE','YCBCR'):
                        profile.pop(key, None)
                    with rasterio.open(collar_path, 'w', **profile) as dst:
                        dst.write(clipped_data)
                    ds.close()
                    ds = rasterio.open(collar_path)
                    collar_paths.append(collar_path)

            # Reproject to match first dataset CRS if needed
            if datasets and ds.crs != datasets[0].crs:
                print(f"    Reprojecting quad {k+1} to match CRS...")
                from rasterio.warp import calculate_default_transform, reproject as warp_reproject
                transform, width, height = calculate_default_transform(
                    ds.crs, datasets[0].crs, ds.width, ds.height, *ds.bounds)
                reproj_path = p.parent / f'_reproj_{k}.tif'
                profile = ds.profile.copy()
                profile.update({
                    'driver': 'GTiff', 'crs': datasets[0].crs,
                    'transform': transform, 'width': width, 'height': height,
                    'compress': 'deflate',
                    'photometric': 'RGB' if ds.count >= 3 else 'MINISBLACK',
                })
                for key in ('COMPRESS','PHOTOMETRIC','PREDICTOR','TILED',
                            'BLOCKXSIZE','BLOCKYSIZE','INTERLEAVE','YCBCR'):
                    profile.pop(key, None)
                with rasterio.open(reproj_path, 'w', **profile) as dst:
                    for band in range(1, ds.count + 1):
                        warp_reproject(
                            source=rasterio.band(ds, band),
                            destination=rasterio.band(dst, band),
                            src_transform=ds.transform, src_crs=ds.crs,
                            dst_transform=transform, dst_crs=datasets[0].crs,
                            resampling=Resampling.bilinear)
                ds.close()
                ds = rasterio.open(reproj_path)
                reproj_paths.append(reproj_path)

            datasets.append(ds)

        if not datasets:
            return False
        if len(datasets) == 1:
            result = _clip_and_save(datasets[0], None, site_geom, clip_wgs84, out_jpg)
            return result

        # Merge collar-clipped quads
        mosaic_data, mosaic_transform = rasterio_merge(datasets, method='first', nodata=255)

        profile = datasets[0].profile.copy()
        profile.update({
            'driver': 'GTiff', 'compress': 'deflate',
            'photometric': 'RGB' if mosaic_data.shape[0] >= 3 else 'MINISBLACK',
            'height': mosaic_data.shape[1], 'width': mosaic_data.shape[2],
            'transform': mosaic_transform, 'count': mosaic_data.shape[0],
            'dtype': mosaic_data.dtype,
        })
        for key in ('COMPRESS','PHOTOMETRIC','PREDICTOR','TILED',
                    'BLOCKXSIZE','BLOCKYSIZE','INTERLEAVE','YCBCR'):
            profile.pop(key, None)
        with rasterio.open(tmp_mosaic, 'w', **profile) as dst:
            dst.write(mosaic_data)

        for ds in datasets:
            try: ds.close()
            except: pass
        datasets = []

        with rasterio.open(tmp_mosaic) as src:
            result = _clip_and_save(src, None, site_geom, clip_wgs84, out_jpg)
        return result

    except Exception as e:
        print(f"    Mosaic error: {e}")
        import traceback; traceback.print_exc()
        if datasets:
            try:
                print("    Falling back to first quad only...")
                result = _clip_and_save(datasets[0], None, site_geom, clip_wgs84, out_jpg)
                return result
            except: pass
        return False
    finally:
        for ds in datasets:
            try: ds.close()
            except: pass
        for p in [tmp_mosaic] + reproj_paths + collar_paths:
            try: p.unlink(missing_ok=True)
            except: pass



def _clip_and_save(src, _, site_geom, clip_wgs84, out_jpg):
    rcrs = src.crs
    t    = Transformer.from_crs('EPSG:4326', rcrs, always_xy=True)
    c    = clip_wgs84
    x0,y0 = t.transform(c[0], c[1])
    x1,y1 = t.transform(c[2], c[3])
    cb   = (min(x0,x1), min(y0,y1), max(x0,x1), max(y0,y1))
    win  = win_from_bounds(*cb, transform=src.transform)
    win  = win.intersection(rasterio.windows.Window(0, 0, src.width, src.height))
    if win.width <= 0 or win.height <= 0:
        print("    SKIP: outside raster"); return False
    oh = min(max(int(win.height), 1), src.height)
    ow = min(max(int(win.width),  1), src.width)
    data = src.read(window=win, out_shape=(src.count, oh, ow),
                    resampling=Resampling.bilinear)
    otf  = src.window_transform(win)

    if data.ndim == 3 and data.shape[0] == 1:
        b   = data[0].astype(np.uint8)
        rgb = np.stack([b, b, b], axis=0)
    elif data.ndim == 3 and data.shape[0] >= 3:
        rgb = data[:3].astype(np.uint8)
    else:
        return False

    img  = Image.fromarray(np.transpose(rgb, (1,2,0)), 'RGB')
    w, h = img.size
    draw = ImageDraw.Draw(img)
    proj = Transformer.from_crs('EPSG:4326', rcrs, always_xy=True)

    def px(lng, lat):
        rx, ry = proj.transform(lng, lat)
        col = (rx - otf.c) / otf.a
        row = (ry - otf.f) / otf.e
        return int(col*w/data.shape[2]), int(row*h/data.shape[1])

    lw    = max(2, w // 250)
    geoms = site_geom.geoms if site_geom.geom_type == 'MultiPolygon' else [site_geom]
    for poly in geoms:
        pts = [px(lng, lat) for lng, lat in poly.exterior.coords]
        draw.line(pts + [pts[0]], fill=SITE_COLOR, width=lw)

    img = img.resize((IMG_W_PX, IMG_H_PX), Image.LANCZOS)
    img.save(out_jpg, 'JPEG', quality=JPEG_QUALITY, optimize=True)
    kb = Path(out_jpg).stat().st_size // 1024
    print(f"    Rendered: {kb} KB")
    return True


# ── Word helpers ──────────────────────────────────────────────────────────────
def get_or_add(el, tag):
    child = el.find(qn(tag))
    if child is None:
        child = OxmlElement(tag)
        el.insert(0, child)
    return child


def tbl_border(tbl, val='single', sz=12, color='000000'):
    tblPr = get_or_add(tbl._tbl, 'w:tblPr')
    b = OxmlElement('w:tblBorders')
    for e in ('top','left','bottom','right','insideH','insideV'):
        el = OxmlElement(f'w:{e}')
        el.set(qn('w:val'), val); el.set(qn('w:sz'), str(sz))
        el.set(qn('w:space'), '0'); el.set(qn('w:color'), color)
        b.append(el)
    tblPr.append(b)


def no_border(cell):
    tcPr = get_or_add(cell._tc, 'w:tcPr')
    b = OxmlElement('w:tcBorders')
    for e in ('top','left','bottom','right','insideH','insideV'):
        el = OxmlElement(f'w:{e}')
        el.set(qn('w:val'), 'none'); el.set(qn('w:sz'), '0')
        el.set(qn('w:space'), '0'); el.set(qn('w:color'), 'auto')
        b.append(el)
    tcPr.append(b)


def cell_w(cell, emu):
    tcPr = get_or_add(cell._tc, 'w:tcPr')
    el   = OxmlElement('w:tcW')
    el.set(qn('w:w'), str(int(emu / 914400 * 1440)))
    el.set(qn('w:type'), 'dxa')
    tcPr.append(el)


def row_h(row, emu, exact=True):
    trPr = get_or_add(row._tr, 'w:trPr')
    el   = OxmlElement('w:trHeight')
    el.set(qn('w:val'), str(int(emu / 914400 * 1440)))
    if exact: el.set(qn('w:hRule'), 'exact')
    trPr.append(el)


def run(para, text, bold=False, italic=False, pt=11, font='Segoe UI', color=None):
    r = para.add_run(text)
    r.bold = bold; r.italic = italic
    r.font.name = font; r.font.size = Pt(pt)
    if color: r.font.color.rgb = RGBColor(*color)
    return r


# ── Document builder ──────────────────────────────────────────────────────────
def build_doc(pages, project_no, out_path, north_arrow):
    doc = Document()
    for sec in doc.sections:
        sec.page_width     = PAGE_W;  sec.page_height   = PAGE_H
        sec.left_margin    = MAR_L;   sec.right_margin  = MAR_R
        sec.top_margin     = MAR_TOP; sec.bottom_margin = MAR_BOT
        sec.header_distance = Inches(0); sec.footer_distance = Inches(0)
    for p in doc.paragraphs:
        p._element.getparent().remove(p._element)

    cw    = PAGE_W - MAR_L - MAR_R  # 7.55"
    cap_w = Inches(5.3)              # left caption column
    key_w = Inches(2.25)             # right KEY column

    for i, page in enumerate(pages):
        meta = page['meta']
        img  = page['img']
        if i > 0: doc.add_page_break()

        # ── Blank-page fix ────────────────────────────────────────────────
        # A mosaic title can wrap to 2-3 lines, pushing the caption past the
        # page bottom and creating a blank page. Estimate the wrapped title
        # height and shrink the map image just enough to keep one page.
        title_str = f"USGS {meta['scale']}  {meta['quad']}, {meta['state']} Quadrangle"
        # ~74 chars fit on the 5.3" caption line at 11pt Segoe UI
        CHARS_PER_LINE = 74
        title_lines    = max(1, -(-len(title_str) // CHARS_PER_LINE))  # ceil division
        # 1 title line is the baseline; each extra line steals ~0.21" — give
        # that back by trimming the map height (which is otherwise 8.5").
        extra_lines    = max(0, title_lines - 1)
        map_h          = MAP_IMG_H - Inches(0.24 * extra_lines)
        # Caption font: 11pt normally, 10pt when wrapping, to claw back room
        cap_pt         = 10 if extra_lines else 11

        # ── Map image in framed table ─────────────────────────────────────
        # Use autofit so the cell exactly wraps the image — no separate height
        mt = doc.add_table(1, 1)
        tbl_border(mt)
        mt.alignment = WD_TABLE_ALIGNMENT.LEFT

        # Set table to fixed width = content width
        tblPr = get_or_add(mt._tbl, 'w:tblPr')
        tblW  = OxmlElement('w:tblW')
        tblW.set(qn('w:w'), str(int(cw / 914400 * 1440)))
        tblW.set(qn('w:type'), 'dxa')
        tblPr.append(tblW)

        mc = mt.cell(0, 0)
        cell_w(mc, cw)
        # Remove cell padding so image sits flush against border
        tcPr  = get_or_add(mc._tc, 'w:tcPr')
        tcMar = OxmlElement('w:tcMar')
        for edge in ('top','left','bottom','right'):
            el = OxmlElement(f'w:{edge}')
            el.set(qn('w:w'), '0'); el.set(qn('w:type'), 'dxa')
            tcMar.append(el)
        tcPr.append(tcMar)

        mp = mc.paragraphs[0]
        mp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        mp.paragraph_format.space_before = Pt(0)
        mp.paragraph_format.space_after  = Pt(0)
        if Path(img).exists():
            # Set image to exact content width, height trimmed for wrapped titles
            mp.add_run().add_picture(img, width=MAP_IMG_W, height=map_h)

        # ── Caption row ───────────────────────────────────────────────────
        ct  = doc.add_table(1, 2)
        ct.alignment = WD_TABLE_ALIGNMENT.LEFT
        lc = ct.cell(0, 0); rc = ct.cell(0, 1)
        cell_w(lc, cap_w); cell_w(rc, key_w)
        no_border(lc); no_border(rc)

        # Left: map title + edition
        lp1 = lc.paragraphs[0]
        lp1.paragraph_format.space_before = Pt(0)
        lp1.paragraph_format.space_after  = Pt(0)
        run(lp1, 'USGS ', pt=cap_pt); run(lp1, meta['scale'] + '  ', pt=cap_pt)
        run(lp1, f"{meta['quad']}, {meta['state']}", italic=True, pt=cap_pt)
        run(lp1, ' Quadrangle', pt=cap_pt)

        lp2 = lc.add_paragraph()
        lp2.paragraph_format.space_before = Pt(0)
        lp2.paragraph_format.space_after  = Pt(0)
        # Edition string e.g. "Created: 1987 (HTMC, 1994 ed.)"
        edition_txt = f"Created: {meta['year']}"
        if meta.get('edition'):
            edition_txt += meta['edition']
        run(lp2, edition_txt, pt=cap_pt)

        # Right: KEY + north arrow + Subject Property
        rp1 = rc.paragraphs[0]
        rp1.paragraph_format.space_before = Pt(0)
        rp1.paragraph_format.space_after  = Pt(0)
        run(rp1, 'KEY:', bold=True, pt=9)
        rp2 = rc.add_paragraph()
        rp2.paragraph_format.space_before = Pt(0)
        rp2.paragraph_format.space_after  = Pt(0)
        if north_arrow and Path(north_arrow).exists():
            rp2.add_run().add_picture(str(north_arrow), height=Inches(0.28))
            run(rp2, '  ', pt=9)
        sym = rp2.add_run('━━  ')
        sym.font.color.rgb = RGBColor(255, 0, 0); sym.font.size = Pt(12)
        run(rp2, 'Subject Property', pt=9)

    # ── Footer ────────────────────────────────────────────────────────────
    # Only set footer on first section; link all others to it
    for sec_idx, sec in enumerate(doc.sections):
        if sec_idx > 0:
            # Link footer to previous section to avoid blank pages
            sec.footer.is_linked_to_previous = True
            continue
        ftr = sec.footer
        for p in ftr.paragraphs:
            p._element.getparent().remove(p._element)
        ft  = ftr.add_table(1, 3, width=Inches(7.55))
        ft.alignment = WD_TABLE_ALIGNMENT.LEFT
        lc2 = ft.cell(0,0); mc2 = ft.cell(0,1); rc2 = ft.cell(0,2)
        cell_w(lc2, Inches(4.0)); cell_w(mc2, Inches(1.5)); cell_w(rc2, Inches(2.05))
        for c in (lc2, mc2, rc2): no_border(c)

        lf  = lc2.paragraphs[0]
        r1  = lf.add_run('Appendix B:  Historic Topographic Maps')
        r1.bold = True; r1.font.name = 'Segoe UI'
        r1.font.size = Pt(14); r1.font.all_caps = True
        lf2 = lc2.add_paragraph()
        r2  = lf2.add_run(f'Project No. {project_no}')
        r2.bold = True; r2.font.name = 'Segoe UI'; r2.font.size = Pt(10)

        rf = rc2.paragraphs[0]
        rf.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        rf.paragraph_format.space_before = Pt(6)
        cr = rf.add_run('Chenier Environmental Consulting, LLC')
        cr.font.name = 'Segoe UI'; cr.font.size = Pt(10)

        # Top border
        tP  = get_or_add(ft._tbl, 'w:tblPr')
        brd = OxmlElement('w:tblBorders')
        tp  = OxmlElement('w:top')
        tp.set(qn('w:val'), 'single'); tp.set(qn('w:sz'), '4')
        tp.set(qn('w:color'), 'auto')
        brd.append(tp); tP.append(brd)

    # out_path may be a filesystem path or a file-like object (BytesIO)
    if hasattr(out_path, 'write'):
        doc.save(out_path)
    else:
        doc.save(str(out_path))
        sz = out_path.stat().st_size
        print(f"\n  Saved: {out_path}")
        print(f"  Size:  {sz//1024} KB  ({sz//1048576} MB)")


# ── Main ──────────────────────────────────────────────────────────────────────